import io
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch


def _parse_achievements(work_entry) -> list:
    """Safely parse achievements from a work experience entry."""
    achievements = []
    try:
        if hasattr(work_entry, 'achievements') and work_entry.achievements:
            achievements = json.loads(work_entry.achievements)
    except:
        pass
    return achievements


def _parse_json_list(user_data, field_name) -> list:
    """Safely parse a JSON list field from user data."""
    result = []
    try:
        val = getattr(user_data, field_name, None)
        if val:
            result = json.loads(val)
    except:
        pass
    return result


def _get_date_range(entry) -> str:
    """Build a clean date range string from an entry."""
    if hasattr(entry, 'start_date') and entry.start_date:
        end = getattr(entry, 'end_date', None) or "Present"
        return f"{entry.start_date} — {end}"
    if hasattr(entry, 'period') and entry.period:
        return entry.period
    if hasattr(entry, 'graduation_year') and entry.graduation_year:
        return entry.graduation_year
    return ""


# ──────────────────────────────────────────────────────────────────────────────
# DOCX GENERATION — Premium Executive Layout
# ──────────────────────────────────────────────────────────────────────────────

def generate_docx_resume(user_data, template="ats") -> io.BytesIO:
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Typography config per template
    is_modern = template in ("modern", "europass")
    heading_font = "Calibri" if is_modern else "Georgia"
    body_font = "Calibri" if is_modern else "Georgia"
    name_color = RGBColor(23, 37, 84) if is_modern else RGBColor(0, 0, 0)
    heading_color = RGBColor(23, 37, 84) if is_modern else RGBColor(30, 41, 59)
    body_color = RGBColor(51, 65, 85)
    muted_color = RGBColor(100, 116, 139)
    divider_color = RGBColor(203, 213, 225) if is_modern else RGBColor(148, 163, 184)

    # Normal style baseline
    style = doc.styles['Normal']
    style.font.name = body_font
    style.font.size = Pt(10)
    style.font.color.rgb = body_color
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.15

    # ── NAME ──
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_after = Pt(2)
    name_run = name_p.add_run(user_data.full_name.upper())
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.name = heading_font
    name_run.font.color.rgb = name_color
    name_run.font.letter_spacing = Pt(2.5)

    # ── CONTACT LINE ──
    contact_parts = []
    if user_data.email:
        contact_parts.append(user_data.email)
    if user_data.phone:
        contact_parts.append(user_data.phone)
    loc = ", ".join(filter(None, [
        getattr(user_data, 'city', None),
        getattr(user_data, 'country', None)
    ]))
    if loc:
        contact_parts.append(loc)
    if user_data.linkedin:
        contact_parts.append(user_data.linkedin)

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.paragraph_format.space_after = Pt(8)
        contact_run = contact_p.add_run("  ·  ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.name = body_font
        contact_run.font.color.rgb = muted_color

    # ── SECTION HEADING HELPER ──
    def add_section_heading(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = heading_font
        run.font.color.rgb = heading_color
        run.font.letter_spacing = Pt(1.5)

        # Thin horizontal rule
        hr_p = doc.add_paragraph()
        hr_p.paragraph_format.space_before = Pt(2)
        hr_p.paragraph_format.space_after = Pt(6)
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="{divider_color.hex()}"/>'
            f'</w:pBdr>'
        )
        hr_p._p.get_or_add_pPr().append(pBdr)

    # ── PROFESSIONAL SUMMARY ──
    if user_data.summary:
        add_section_heading("Professional Summary")
        summary_p = doc.add_paragraph()
        summary_p.paragraph_format.space_after = Pt(4)
        summary_run = summary_p.add_run(user_data.summary)
        summary_run.font.size = Pt(10)
        summary_run.font.color.rgb = body_color
        summary_run.italic = True

    # ── WORK EXPERIENCE ──
    if user_data.work_experience:
        add_section_heading("Professional Experience")
        for exp in user_data.work_experience:
            # Role + Company line
            role_p = doc.add_paragraph()
            role_p.paragraph_format.space_before = Pt(6)
            role_p.paragraph_format.space_after = Pt(1)
            role_run = role_p.add_run(exp.role)
            role_run.bold = True
            role_run.font.size = Pt(10.5)
            role_run.font.color.rgb = RGBColor(15, 23, 42)
            role_p.add_run("  —  ").font.color.rgb = muted_color
            company_run = role_p.add_run(exp.company)
            company_run.font.color.rgb = body_color

            # Date range
            date_str = _get_date_range(exp)
            if date_str:
                date_p = doc.add_paragraph()
                date_p.paragraph_format.space_after = Pt(2)
                date_run = date_p.add_run(date_str)
                date_run.italic = True
                date_run.font.size = Pt(9)
                date_run.font.color.rgb = muted_color

            # Achievements as bullet points
            achievements = _parse_achievements(exp)
            if achievements:
                for ach in achievements:
                    bullet_p = doc.add_paragraph()
                    bullet_p.paragraph_format.left_indent = Inches(0.25)
                    bullet_p.paragraph_format.space_after = Pt(1)
                    bullet_run = bullet_p.add_run(f"▸  {ach}")
                    bullet_run.font.size = Pt(9.5)
                    bullet_run.font.color.rgb = body_color
            elif exp.description:
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.left_indent = Inches(0.25)
                desc_p.paragraph_format.space_after = Pt(2)
                desc_run = desc_p.add_run(exp.description)
                desc_run.font.size = Pt(9.5)
                desc_run.font.color.rgb = body_color

    # ── EDUCATION ──
    if user_data.education:
        add_section_heading("Education")
        for edu in user_data.education:
            # Degree + School
            edu_p = doc.add_paragraph()
            edu_p.paragraph_format.space_before = Pt(4)
            edu_p.paragraph_format.space_after = Pt(1)

            degree_text = edu.degree
            if hasattr(edu, 'major') and edu.major:
                degree_text += f" in {edu.major}"

            deg_run = edu_p.add_run(degree_text)
            deg_run.bold = True
            deg_run.font.size = Pt(10.5)
            deg_run.font.color.rgb = RGBColor(15, 23, 42)

            edu_p.add_run("  —  ").font.color.rgb = muted_color
            school_run = edu_p.add_run(edu.school)
            school_run.font.color.rgb = body_color

            # Period / Graduation Year + GPA
            details_parts = []
            date_str = _get_date_range(edu)
            if date_str:
                details_parts.append(date_str)
            if edu.gpa:
                details_parts.append(f"GPA: {edu.gpa}")

            if details_parts:
                detail_p = doc.add_paragraph()
                detail_p.paragraph_format.space_after = Pt(2)
                detail_run = detail_p.add_run("  ·  ".join(details_parts))
                detail_run.italic = True
                detail_run.font.size = Pt(9)
                detail_run.font.color.rgb = muted_color

            if hasattr(edu, 'details') and edu.details:
                det_p = doc.add_paragraph()
                det_p.paragraph_format.left_indent = Inches(0.25)
                det_p.paragraph_format.space_after = Pt(2)
                det_run = det_p.add_run(edu.details)
                det_run.font.size = Pt(9.5)
                det_run.font.color.rgb = body_color

    # ── PROJECTS & CERTIFICATIONS ──
    if user_data.projects:
        add_section_heading("Projects & Certifications")
        for proj in user_data.projects:
            proj_p = doc.add_paragraph()
            proj_p.paragraph_format.space_before = Pt(4)
            proj_p.paragraph_format.space_after = Pt(1)
            proj_run = proj_p.add_run(proj.name)
            proj_run.bold = True
            proj_run.font.size = Pt(10)
            proj_run.font.color.rgb = RGBColor(15, 23, 42)

            if proj.description:
                desc_p = doc.add_paragraph()
                desc_p.paragraph_format.left_indent = Inches(0.25)
                desc_p.paragraph_format.space_after = Pt(1)
                desc_run = desc_p.add_run(proj.description)
                desc_run.font.size = Pt(9.5)
                desc_run.font.color.rgb = body_color

            if hasattr(proj, 'link_or_credential') and proj.link_or_credential:
                link_p = doc.add_paragraph()
                link_p.paragraph_format.left_indent = Inches(0.25)
                link_p.paragraph_format.space_after = Pt(2)
                link_run = link_p.add_run(proj.link_or_credential)
                link_run.font.size = Pt(8.5)
                link_run.font.color.rgb = RGBColor(37, 99, 235)
                link_run.italic = True

    # ── SKILLS ──
    skills_list = _parse_json_list(user_data, 'skills')
    languages_list = _parse_json_list(user_data, 'languages')

    if skills_list or languages_list:
        add_section_heading("Skills & Languages")
        if skills_list:
            tech_p = doc.add_paragraph()
            tech_p.paragraph_format.space_after = Pt(2)
            label_run = tech_p.add_run("Technical:  ")
            label_run.bold = True
            label_run.font.size = Pt(9.5)
            label_run.font.color.rgb = body_color
            val_run = tech_p.add_run("  ·  ".join(skills_list))
            val_run.font.size = Pt(9.5)
            val_run.font.color.rgb = body_color

        if languages_list:
            lang_p = doc.add_paragraph()
            lang_p.paragraph_format.space_after = Pt(2)
            label_run = lang_p.add_run("Languages:  ")
            label_run.bold = True
            label_run.font.size = Pt(9.5)
            label_run.font.color.rgb = body_color
            val_run = lang_p.add_run("  ·  ".join(languages_list))
            val_run.font.size = Pt(9.5)
            val_run.font.color.rgb = body_color

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


# ──────────────────────────────────────────────────────────────────────────────
# PDF GENERATION — Premium Executive Layout
# ──────────────────────────────────────────────────────────────────────────────

def generate_pdf_resume(user_data, template="ats") -> io.BytesIO:
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=48,
        leftMargin=48,
        topMargin=48,
        bottomMargin=48
    )

    is_modern = template in ("modern", "europass")

    # ── COLOR PALETTE ──
    dark_navy = colors.HexColor("#172554")
    slate_700 = colors.HexColor("#334155")
    slate_500 = colors.HexColor("#64748b")
    slate_300 = colors.HexColor("#cbd5e1")
    accent_blue = colors.HexColor("#1e40af") if is_modern else dark_navy
    body_text_color = slate_700 if is_modern else colors.HexColor("#1e293b")

    # ── TYPOGRAPHY STYLES ──
    base_font = "Helvetica" if is_modern else "Times-Roman"
    bold_font = "Helvetica-Bold" if is_modern else "Times-Bold"
    italic_font = "Helvetica-Oblique" if is_modern else "Times-Italic"

    name_style = ParagraphStyle(
        'CVName',
        fontName=bold_font,
        fontSize=22,
        leading=26,
        alignment=1,
        textColor=dark_navy,
        spaceAfter=2,
        tracking=200,
    )

    contact_style = ParagraphStyle(
        'CVContact',
        fontName=base_font,
        fontSize=9,
        leading=13,
        alignment=1,
        textColor=slate_500,
        spaceAfter=10,
    )

    section_heading_style = ParagraphStyle(
        'CVSectionHeading',
        fontName=bold_font,
        fontSize=10,
        leading=14,
        textColor=accent_blue,
        spaceBefore=16,
        spaceAfter=2,
        tracking=150,
    )

    body_style = ParagraphStyle(
        'CVBody',
        fontName=base_font,
        fontSize=10,
        leading=14,
        textColor=body_text_color,
    )

    summary_style = ParagraphStyle(
        'CVSummary',
        fontName=italic_font,
        fontSize=10,
        leading=14.5,
        textColor=body_text_color,
        spaceAfter=4,
    )

    role_style = ParagraphStyle(
        'CVRole',
        fontName=base_font,
        fontSize=10,
        leading=14,
        textColor=body_text_color,
        spaceBefore=6,
        spaceAfter=1,
    )

    date_style = ParagraphStyle(
        'CVDate',
        fontName=italic_font,
        fontSize=9,
        leading=12,
        textColor=slate_500,
        alignment=2,
    )

    bullet_style = ParagraphStyle(
        'CVBullet',
        fontName=base_font,
        fontSize=9.5,
        leading=13.5,
        textColor=body_text_color,
        leftIndent=14,
        spaceAfter=1,
    )

    link_style = ParagraphStyle(
        'CVLink',
        fontName=italic_font,
        fontSize=8.5,
        leading=12,
        textColor=colors.HexColor("#2563eb"),
        leftIndent=14,
        spaceAfter=2,
    )

    label_style = ParagraphStyle(
        'CVLabel',
        fontName=bold_font,
        fontSize=9.5,
        leading=13,
        textColor=body_text_color,
    )

    story = []

    # ── NAME ──
    story.append(Paragraph(user_data.full_name.upper(), name_style))
    story.append(Spacer(1, 2))

    # ── CONTACT INFO ──
    contact_parts = []
    if user_data.email:
        contact_parts.append(user_data.email)
    if user_data.phone:
        contact_parts.append(user_data.phone)
    loc = ", ".join(filter(None, [
        getattr(user_data, 'city', None),
        getattr(user_data, 'country', None)
    ]))
    if loc:
        contact_parts.append(loc)
    if user_data.linkedin:
        contact_parts.append(user_data.linkedin)

    if contact_parts:
        story.append(Paragraph("  ·  ".join(contact_parts), contact_style))

    # ── DIVIDER ──
    story.append(HRFlowable(
        width="100%", thickness=0.8,
        color=accent_blue if is_modern else slate_300,
        spaceBefore=2, spaceAfter=6
    ))

    # ── SECTION HELPER ──
    def add_section(title):
        story.append(Paragraph(title.upper(), section_heading_style))
        story.append(HRFlowable(
            width="100%", thickness=0.5,
            color=slate_300,
            spaceBefore=1, spaceAfter=6
        ))

    usable_width = letter[0] - 96  # 48pt margins each side

    # ── PROFESSIONAL SUMMARY ──
    if user_data.summary:
        add_section("Professional Summary")
        story.append(Paragraph(user_data.summary, summary_style))
        story.append(Spacer(1, 4))

    # ── WORK EXPERIENCE ──
    if user_data.work_experience:
        add_section("Professional Experience")
        for exp in user_data.work_experience:
            # Role / Company on left — Date on right
            left_text = f"<b>{exp.role}</b>  —  {exp.company}"
            date_text = f"<i>{_get_date_range(exp)}</i>"

            left_p = Paragraph(left_text, role_style)
            right_p = Paragraph(date_text, date_style)

            header_table = Table([[left_p, right_p]], colWidths=[usable_width * 0.68, usable_width * 0.32])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(header_table)

            # Achievement bullets
            achievements = _parse_achievements(exp)
            if achievements:
                for ach in achievements:
                    story.append(Paragraph(f"▸  {ach}", bullet_style))
            elif exp.description:
                story.append(Paragraph(exp.description, bullet_style))

            story.append(Spacer(1, 6))

    # ── EDUCATION ──
    if user_data.education:
        add_section("Education")
        for edu in user_data.education:
            degree_text = edu.degree
            if hasattr(edu, 'major') and edu.major:
                degree_text += f" in {edu.major}"

            left_text = f"<b>{degree_text}</b>  —  {edu.school}"
            date_text = f"<i>{_get_date_range(edu)}</i>"

            left_p = Paragraph(left_text, role_style)
            right_p = Paragraph(date_text, date_style)

            header_table = Table([[left_p, right_p]], colWidths=[usable_width * 0.68, usable_width * 0.32])
            header_table.setStyle(TableStyle([
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 0),
                ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                ('TOPPADDING', (0, 0), (-1, -1), 0),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
            ]))
            story.append(header_table)

            # GPA + details
            details_parts = []
            if edu.gpa:
                details_parts.append(f"GPA: {edu.gpa}")
            if hasattr(edu, 'details') and edu.details:
                details_parts.append(edu.details)
            if details_parts:
                story.append(Paragraph("  ·  ".join(details_parts), bullet_style))

            story.append(Spacer(1, 6))

    # ── PROJECTS & CERTIFICATIONS ──
    if user_data.projects:
        add_section("Projects & Certifications")
        for proj in user_data.projects:
            story.append(Paragraph(f"<b>{proj.name}</b>", role_style))
            if proj.description:
                story.append(Paragraph(proj.description, bullet_style))
            if hasattr(proj, 'link_or_credential') and proj.link_or_credential:
                story.append(Paragraph(proj.link_or_credential, link_style))
            story.append(Spacer(1, 4))

    # ── SKILLS & LANGUAGES ──
    skills_list = _parse_json_list(user_data, 'skills')
    languages_list = _parse_json_list(user_data, 'languages')

    if skills_list or languages_list:
        add_section("Skills & Languages")
        if skills_list:
            story.append(Paragraph(
                f"<b>Technical:</b>  {'  ·  '.join(skills_list)}",
                ParagraphStyle('SkillLine', parent=body_style, fontSize=9.5, leading=13, spaceAfter=3)
            ))
        if languages_list:
            story.append(Paragraph(
                f"<b>Languages:</b>  {'  ·  '.join(languages_list)}",
                ParagraphStyle('LangLine', parent=body_style, fontSize=9.5, leading=13, spaceAfter=3)
            ))

    # ── BUILD ──
    doc.build(story)
    buffer.seek(0)
    return buffer
